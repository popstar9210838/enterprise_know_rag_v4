"""
文档加载器 —— 自定义 Reader，支持 PDF / DOCX / Excel。
"""
import logging
from pathlib import Path
from typing import List

import pdfplumber

logger = logging.getLogger(__name__)
from llama_index.core import Document
from llama_index.core.readers.base import BaseReader


class PDFPlumberReader(BaseReader):
    """用 pdfplumber 提取 PDF 文本，逐页为独立 Document。"""

    def load_data(self, file_path: str, extra_info: dict = None) -> List[Document]:
        docs = []
        file_name = Path(file_path).name
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        docs.append(Document(
                            text=text.strip(),
                            metadata={"file_name": file_name, "source": file_name, "page": i + 1},
                        ))
        except Exception as e:
            logger.warning("PDF 解析失败: %s (%s)", file_path, e)
            return []
        return docs


class DocxReader(BaseReader):
    """用 python-docx 提取 Word 文本，整个文档为一个 Document。"""

    def load_data(self, file_path: str, extra_info: dict = None) -> List[Document]:
        file_name = Path(file_path).name
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                return []
            return [Document(
                text="\n".join(paragraphs),
                metadata={"file_name": file_name, "source": file_name},
            )]
        except Exception as e:
            logger.warning("DOCX 解析失败: %s (%s)", file_path, e)
            return []


class ExcelReader(BaseReader):
    """用 openpyxl 提取 Excel 文本，每个 Sheet 为一个 Document。"""

    def load_data(self, file_path: str, extra_info: dict = None) -> List[Document]:
        docs = []
        file_name = Path(file_path).name
        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines = []
                for row in ws.iter_rows(values_only=True):
                    values = [str(v) for v in row if v is not None]
                    if values:
                        lines.append(" | ".join(values))
                if lines:
                    docs.append(Document(
                        text="\n".join(lines),
                        metadata={"file_name": file_name, "source": file_name, "sheet": sheet_name},
                    ))
            wb.close()
        except Exception as e:
            logger.warning("Excel 解析失败: %s (%s)", file_path, e)
            return []
        return docs
